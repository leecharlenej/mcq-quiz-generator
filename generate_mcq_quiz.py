# ===========================================================================
# Initialization
# ===========================================================================

import pandas as pd
from docx import Document
from docx.shared import Inches
import random
from collections import Counter

import helper_functions

# ===========================================================================
# Function: Main
# ===========================================================================

def generate_mcq_quiz():

    print('\n-------------- Generate quiz -----------------')
    mcq_quiz_folder_path, mcq_quiz_db_file_path, quiz_df, topics_num_questions = quiz_admin()

    new_quiz_df = randomize_questions_selection(mcq_quiz_folder_path, mcq_quiz_db_file_path, quiz_df, topics_num_questions)

    generate_word_doc(mcq_quiz_folder_path, new_quiz_df)


## ===========================================================================
# Function: Admin to get quiz folder path and number of questions for each topic
# ===========================================================================

def quiz_admin():

    # User to input the name of the quiz folder.
    mcq_quiz_folder_path = helper_functions.user_input("Enter path of quiz folder (E.g. G:\My Drive\quiz_folder): ")

    # Get headers from the quiz database and create a dictionary to store the new question.
    mcq_quiz_db_file_path = mcq_quiz_folder_path + "/quiz_db.csv"
    quiz_df = pd.read_csv(mcq_quiz_db_file_path)
    quiz_topics = quiz_df["Topic"].unique()
    
    # Number of questions for each topic
    topics_num_questions = {}
    for i in quiz_topics:
        topics_num_questions[i] = 0
    
    same_num_per_topic = helper_functions.user_input("Do you want same number of questions for each topic? (Y/N) [Yes/No] ").strip().lower()
    if same_num_per_topic == "y" or same_num_per_topic == "yes":
        num_questions = int(helper_functions.user_input("Enter number of questions for each topic: "))
        for i in quiz_topics:
            topics_num_questions[i] = num_questions
    else:
        for i in quiz_topics:
            num_questions = int(helper_functions.user_input(f"Enter number of questions for {i}: "))
            topics_num_questions[i] = num_questions

    print(topics_num_questions)
    
    return (mcq_quiz_folder_path, mcq_quiz_db_file_path, quiz_df, topics_num_questions)

## ===========================================================================
# Function: Randomize questions selection + include ChatGPT questions?
# ===========================================================================
    
def randomize_questions_selection(mcq_quiz_folder_path, mcq_quiz_db_file_path, quiz_df, topics_num_questions):

    # Load quiz_db.csv to get headers and create new DataFrame for selected quiz questions.
    quiz_df = pd.read_csv(mcq_quiz_db_file_path, keep_default_na=False)
    quiz_headers = quiz_df.columns
    
    new_quiz_df = pd.DataFrame(columns=quiz_headers)

    include_chatgpt = helper_functions.user_input("Do you want to include ChatGPT questions? (Y/N) [Yes/No] ").strip().lower()

    if include_chatgpt == "n" or include_chatgpt == "no":
        
        # No ChatGPT
        quiz_df_noChatGPT = quiz_df[quiz_df["Type"] == "Base"].copy()

        for topic, num_questions in topics_num_questions.items():

            # Filter the DataFrame for the current topic
            selected_topic_df = quiz_df_noChatGPT[quiz_df_noChatGPT["Topic"] == topic].copy()

            # Check if there are enough questions in the topic
            if len(selected_topic_df) >= num_questions:
            # If enough questions, sample without replacement
                sampled_df = selected_topic_df.sample(n=num_questions)
            else:
                sampled_df = selected_topic_df.sample(n=len(selected_topic_df))

            # Concatenate the sampled questions into the new_quiz_df
            new_quiz_df = pd.concat([new_quiz_df, sampled_df], ignore_index=True)
        
    elif include_chatgpt == "y" or include_chatgpt == "yes":
        
        # ChatGPT
    
        for topic, num_questions in topics_num_questions.items():

            # Filter the DataFrame for the current topic, get the IDs and count of each ID
            selected_topic_df = quiz_df[quiz_df["Topic"] == topic].copy()
            selected_topic_ids = selected_topic_df["ID"].tolist()
            selected_topic_ids_count = Counter(selected_topic_ids)
            selected_topic_ids_list = list(selected_topic_ids_count.keys())

            # Randomly select num_questions of IDs from the list of IDs for current topic
            selected_ids = random.sample(selected_topic_ids_list, num_questions)

            for i in selected_ids:
                # print(i, selected_topic_ids_count[i])
                
                if selected_topic_ids_count[i] > 1:
                    to_add = selected_topic_df[selected_topic_df["ID"] == i].sample(n=1)
                    
                elif selected_topic_ids_count[i] == 1:
                    to_add = selected_topic_df[selected_topic_df["ID"] == i]
    
                new_quiz_df = pd.concat([new_quiz_df, to_add], ignore_index=True)

    return new_quiz_df

## ===========================================================================
# Function: Generate Word document
# ===========================================================================

def generate_word_doc(mcq_quiz_folder_path, new_quiz_df):

    def set_col_widths(table):
        for row in table.rows:
            for cell, width in zip(row.cells, (Inches(0.5), Inches(1), Inches(7))):
                cell.width = width

    new_quiz_headers = new_quiz_df.columns
    new_quiz_option_count = len([header for header in new_quiz_headers if header.startswith('Option')])

    document = Document()
    quiz_title = helper_functions.user_input("Enter title of quiz: ")
    document.add_heading('Test title: ' + quiz_title, 0)

    shuffle_options = helper_functions.user_input("Do you want to shuffle options? (Y/N) [Yes/No] ").strip().upper()
    
    # Quiz header rows
    table = document.add_table(1,3)
    table.autofit = True
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'No.'
    header_cells[1].text = 'Topic/ Sub-topic'
    header_cells[2].text = 'Question'

    # Add quiz questions
    for row in new_quiz_df.itertuples():
        cells = table.add_row().cells
        cells[0].text = str(row.Index+1)

        cells[1].text = f'{row.Topic}/'
        cells[1].add_paragraph(row.Sub_topic)
        cells[1].add_paragraph('\n')
        cells[1].add_paragraph(f'({row.Type})')


        # if row.Image == 'Y' or row.Image == 'YES':
        #     image_path = mcq_quiz_folder_path + '/images/' + row.Image_file
        #     cells[2].paragraphs[0].add_run().add_picture(image_path, width=Inches(1.25))


        cells[2].text = row.Question + '\n'

        option_num = list(range(1, new_quiz_option_count+1))

        if shuffle_options == 'Y' or shuffle_options == 'YES':
            random.shuffle(option_num)
            
        for i in option_num:
                option = getattr(row, f'Option{i}')
                cells[2].add_paragraph(option, style='List Bullet')
        
        cells[2].add_paragraph('\n')

        cells[2].add_paragraph('Answer: ' + row.Answer)


    set = set_col_widths(table)


    document.add_page_break()

    quiz_title_pathName = quiz_title.replace(" ", "_")
    document.save(mcq_quiz_folder_path +'/' + quiz_title_pathName +'.docx')

# ===========================================================================
# Test mode
# ===========================================================================
def test(mode):
    if mode == 1:
        generate_mcq_quiz()

    elif mode == 2:
        mcq_quiz_folder_path = ""
        mcq_quiz_db_file_path = mcq_quiz_folder_path + "/quiz_db.csv"
        quiz_df = pd.read_csv(mcq_quiz_db_file_path)
        topics_num_questions = {'Cryptography': 0, 'Networking and Forensics': 0, 'Digital Forensics': 3}

        df_test = randomize_questions_selection(mcq_quiz_folder_path, mcq_quiz_db_file_path, quiz_df, topics_num_questions)
        print(df_test)


if __name__ == "__main__":
    test(0) # 0 for no test